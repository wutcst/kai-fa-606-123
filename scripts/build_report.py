from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


OUT = "REPORT.docx"
PDF_OUT = "REPORT.pdf"


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=11, bold=False, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def add_paragraph(doc, text="", style=None, size=11, bold=False, color=None, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    color = "2E74B5" if level <= 2 else "1F4D78"
    size = 16 if level == 1 else 13 if level == 2 else 12
    set_run_font(run, size=size, bold=True, color=color)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True, fill="F2F4F7")
        if widths:
            header_cells[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(4)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Starline Valkyrie 横版太空射击游戏\n小组实训报告")
    set_run_font(run, size=22, bold=True, color="0B2545")
    title.paragraph_format.space_after = Pt(24)

    meta_rows = [
        ("课程任务", "软件工程实训任务二：小组协同开发"),
        ("项目名称", "Starline Valkyrie 横版太空射击游戏"),
        ("代码仓库", "https://github.com/wutcst/kai-fa-606-123"),
        ("小组名称", "待补充"),
        ("组长", "待补充"),
        ("成员", "待补充"),
        ("报告日期", "2026-06-20"),
    ]
    add_table(doc, ["项目", "内容"], meta_rows, widths=[1.5, 4.8])

    add_heading(doc, "一、项目概述")
    add_paragraph(
        doc,
        "本项目根据课程任务中“可以不基于 world-of-zuul 样例代码，自选择新的游戏项目内容进行开发”的要求，"
        "完成了一个基于 HTML5 Canvas 的横版街机射击游戏。游戏以即时反馈和低学习成本为核心目标，"
        "玩家打开页面即可移动战机、自动射击、消灭敌机、积累连击并释放 BOMB 清屏技能。",
    )
    add_paragraph(
        doc,
        "项目的主要成果包括完整的网页交互界面、游戏状态机、敌人系统、碰撞系统、分数连击系统、Boss 机制、"
        "程序化视觉素材、粒子爆炸特效、自动化测试和 GitHub Actions 测试流程。"
    )

    add_heading(doc, "二、需求分析")
    add_paragraph(doc, "从原始任务要求出发，项目需要满足三个本质目标：完整可玩、可检查、可提交。")
    add_bullets(
        doc,
        [
            "完整可玩：游戏必须有明确目标、实时交互、胜负反馈、持续挑战和图形界面。",
            "可检查：代码结构应清晰，测试命令可复现，README 和报告能说明设计过程与完成情况。",
            "可提交：所有代码、文档和报告电子版应进入小组 GitHub 仓库，便于教师按仓库工作量评估。",
        ],
    )

    add_heading(doc, "三、总体设计")
    add_paragraph(
        doc,
        "项目采用原生 Web 技术实现，不引入大型框架。这样可以降低部署成本，保证在浏览器中直接运行，"
        "也便于通过 Node.js 内置测试框架验证核心规则。整体结构分为入口层、游戏状态层、规则层和渲染层。"
    )
    add_table(
        doc,
        ["模块", "文件", "职责"],
        [
            ("入口层", "src/main.js", "处理键盘、触控、窗口缩放、HUD 同步和 requestAnimationFrame 主循环。"),
            ("规则层", "src/game/rules.js", "封装自动开火、连击、奖励、火力补给和 Boss 节奏等可测试规则。"),
            ("状态层", "src/game/game.js", "维护玩家、子弹、敌人、掉落、粒子、冲击波、碰撞和游戏状态。"),
            ("渲染层", "src/game/render.js", "负责 Canvas 背景、战机、敌机、子弹、粒子和 UI 反馈绘制。"),
            ("样式层", "src/styles.css", "定义全屏游戏布局、HUD、按钮、暂停和结束面板。"),
            ("测试层", "tests/*.test.js", "使用 node:test 验证规则函数和游戏状态机。"),
        ],
        widths=[1.0, 1.55, 3.75],
    )

    add_heading(doc, "四、关键功能设计与完成情况")
    add_heading(doc, "4.1 玩家与操作", level=2)
    add_paragraph(
        doc,
        "玩家战机位于屏幕左侧，支持 WASD、方向键和触控拖动。射击采用自动开火，玩家不需要持续按键，"
        "从而降低操作负担并突出街机爽感。Space 键或右下角 BOMB 按钮可释放清屏技能。"
    )

    add_heading(doc, "4.2 敌人与 Boss", level=2)
    add_paragraph(
        doc,
        "敌人从右侧进入，包含轻型机、装甲机、追踪机、炮艇和阶段 Boss。普通敌人负责提供连续击杀节奏，"
        "Boss 负责形成阶段性压力和视觉高潮。Boss 会按照时间节奏出现，拥有更高生命值和更密集弹幕。"
    )

    add_heading(doc, "4.3 分数、连击与掉落", level=2)
    add_paragraph(
        doc,
        "击杀敌人会增加分数并刷新连击计时。连击越高，单次击杀得分越高。敌人死亡后可能掉落火力、回血或 BOMB 充能道具，"
        "道具会被玩家磁吸，减少拾取挫败感。"
    )

    add_heading(doc, "4.4 视觉素材与粒子特效", level=2)
    add_paragraph(
        doc,
        "项目没有使用外部贴图库，而是通过 Canvas 程序化绘制战机、敌机、Boss、星空、子弹、冲击波和爆炸粒子。"
        "这种方式避免素材版权问题，也保证整体视觉风格统一。后续针对“更加炸裂”的反馈，新增了双层霓虹冲击波、"
        "高速光束粒子、火花碎片和慢速余烬，让击杀反馈更明显。"
    )

    add_heading(doc, "五、测试与质量保证")
    add_paragraph(doc, "项目使用 Node.js 内置测试框架，执行命令为：npm test。")
    add_table(
        doc,
        ["测试类别", "覆盖内容"],
        [
            ("规则测试", "自动开火、连击衰减、击杀奖励、火力补给上限、Boss 节奏。"),
            ("状态机测试", "游戏循环自动射击和刷怪、BOMB 清弹和伤害、冲击波生命周期。"),
            ("性能边界测试", "子弹、敌弹、掉落物、粒子和冲击波均有数量上限，避免长时间运行后卡顿。"),
            ("浏览器验证", "通过浏览器运行检查 Canvas 非空、无控制台错误、分数和击杀正常增长。"),
        ],
        widths=[1.5, 4.8],
    )
    add_paragraph(
        doc,
        "性能优化过程中，曾发现高 DPR 屏幕和大量 Canvas 阴影/渐变会造成卡顿。优化措施包括背景缓存、"
        "子弹按颜色批量绘制、减少高成本阴影、限制粒子数量、限制 Retina Canvas 最大像素比。"
    )

    add_heading(doc, "六、仓库与协作交付")
    add_paragraph(
        doc,
        "项目已关联小组 GitHub 仓库，并在 README.md 中改写为项目介绍文档。根目录提供 REPORT.docx 作为电子版实训报告。"
        "同时补充 .github/workflows/test.yml，使代码推送或 Pull Request 时可以自动运行 npm test。"
    )
    add_table(
        doc,
        ["课程要求", "当前完成情况"],
        [
            ("在小组仓库中提交开发成果", "已将网页游戏代码、测试、README 和报告加入仓库提交范围。"),
            ("README 作为项目介绍文档", "已重写 README.md，说明项目简介、运行方式、功能、测试和交付信息。"),
            ("根目录 REPORT.docx/REPORT.pdf", "已生成 REPORT.docx；如渲染环境可用，可同步生成 REPORT.pdf。"),
            ("Bilibili 视频展示", "需小组录制 5-10 分钟视频并使用【武理26软工实践】作为标题前缀发布。"),
            ("AI 使用说明", "已在 README 和本报告中说明 AI 模型及辅助工作内容。"),
        ],
        widths=[2.2, 4.1],
    )

    add_heading(doc, "七、AI 辅助开发说明")
    add_paragraph(
        doc,
        "本项目使用 OpenAI Codex / GPT-5 系列能力辅助完成部分设计、开发和文档工作。AI 的作用包括："
    )
    add_bullets(
        doc,
        [
            "辅助梳理课程任务要求，确定可独立实现的横版射击游戏方向。",
            "辅助生成 Canvas 游戏循环、实体系统、碰撞系统、渲染系统和粒子特效代码。",
            "辅助编写自动化测试，覆盖核心规则和游戏状态机。",
            "辅助进行性能诊断，定位 Canvas 渲染瓶颈并提出优化方案。",
            "辅助撰写 README.md 和 REPORT.docx 初稿。",
        ],
    )
    add_paragraph(
        doc,
        "AI 输出并不替代小组审查。最终提交前，小组成员应实际运行游戏、检查报告内容、补充成员分工、"
        "录制演示视频，并在答辩中说明个人完成的真实工作。"
    )

    add_heading(doc, "八、小组分工记录")
    add_paragraph(doc, "以下分工信息需要小组按真实情况补充，不应由 AI 伪造。")
    add_table(
        doc,
        ["成员", "学号", "主要工作", "对应仓库证据"],
        [
            ("待补充", "待补充", "玩法设计 / 核心逻辑 / 测试 / 文档等按实际填写", "提交记录、代码文件、Issue、PR"),
            ("待补充", "待补充", "待补充", "待补充"),
            ("待补充", "待补充", "待补充", "待补充"),
            ("待补充", "待补充", "待补充", "待补充"),
        ],
        widths=[1.0, 1.0, 2.6, 1.7],
    )

    add_heading(doc, "九、总结与后续工作")
    add_paragraph(
        doc,
        "当前项目已经形成一个可运行、可测试、可展示的横版射击游戏。相比原始文本样例工程，项目具备完整图形界面、"
        "实时交互、游戏反馈和自动化测试。后续可继续补充排行榜、更多 Boss、音效、关卡配置、GitHub Issue 分工记录和演示视频链接。"
    )

    doc.save(OUT)
    build_pdf()
    print(OUT)
    print(PDF_OUT)


def p(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf():
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "ChineseBody",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=16,
        spaceAfter=7,
    )
    h1 = ParagraphStyle(
        "ChineseH1",
        parent=styles["Heading1"],
        fontName="STSong-Light",
        fontSize=16,
        leading=22,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=14,
        spaceAfter=8,
    )
    h2 = ParagraphStyle(
        "ChineseH2",
        parent=styles["Heading2"],
        fontName="STSong-Light",
        fontSize=13,
        leading=18,
        textColor=colors.HexColor("#1F4D78"),
        spaceBefore=10,
        spaceAfter=6,
    )
    title_style = ParagraphStyle(
        "ChineseTitle",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=22,
        leading=30,
        alignment=1,
        textColor=colors.HexColor("#0B2545"),
        spaceAfter=20,
    )

    story = [Paragraph("Starline Valkyrie 横版太空射击游戏<br/>小组实训报告", title_style)]
    story.append(make_pdf_table([["项目", "内容"], ["课程任务", "软件工程实训任务二：小组协同开发"], ["项目名称", "Starline Valkyrie 横版太空射击游戏"], ["代码仓库", "https://github.com/wutcst/kai-fa-606-123"], ["小组名称", "待补充"], ["组长", "待补充"], ["成员", "待补充"], ["报告日期", "2026-06-20"]], body, [1.35 * inch, 4.8 * inch]))
    story.append(Spacer(1, 0.15 * inch))

    sections = [
        ("一、项目概述", [
            "本项目根据课程任务中“可以不基于 world-of-zuul 样例代码，自选择新的游戏项目内容进行开发”的要求，完成了一个基于 HTML5 Canvas 的横版街机射击游戏。",
            "项目成果包括完整网页交互界面、游戏状态机、敌人系统、碰撞系统、分数连击系统、Boss 机制、程序化视觉素材、粒子爆炸特效、自动化测试和 GitHub Actions 测试流程。",
        ]),
        ("二、需求分析", [
            "从原始任务要求出发，项目需要满足完整可玩、可检查、可提交三个目标。完整可玩要求游戏具备图形界面、实时交互、胜负反馈和持续挑战；可检查要求代码结构清晰、测试命令可复现；可提交要求代码、文档和报告电子版进入小组 GitHub 仓库。",
        ]),
        ("三、总体设计", [
            "项目采用原生 Web 技术实现，分为入口层、规则层、状态层、渲染层和测试层。rules.js 保存可测试规则；game.js 维护实体和状态；render.js 负责 Canvas 绘制；main.js 负责输入、HUD 和主循环。",
        ]),
        ("四、关键功能设计与完成情况", [
            "玩家支持键盘和触控移动，射击采用自动开火与轻微自动索敌，降低操作负担。敌人包含轻型机、装甲机、追踪机、炮艇和 Boss。击杀会增加分数并刷新连击，掉落物提供火力、回血和 BOMB 充能。",
            "视觉素材通过 Canvas 程序化绘制，不依赖外部贴图。爆炸反馈包含双层霓虹冲击波、高速光束粒子、火花碎片和慢速余烬。",
        ]),
        ("五、测试与质量保证", [
            "项目使用 Node.js 内置 node:test，执行命令为 npm test。测试覆盖自动开火、连击衰减、击杀奖励、火力补给、Boss 节奏、BOMB、冲击波生命周期和实体数量上限。",
            "性能优化包括背景缓存、批量绘制、减少高成本阴影、限制粒子数量、限制 Retina Canvas 最大像素比。",
        ]),
        ("六、仓库与协作交付", [
            "项目已关联 GitHub 仓库并改写 README.md。根目录提供 REPORT.docx 和 REPORT.pdf 作为电子版实训报告。补充 .github/workflows/test.yml，在推送或 Pull Request 时运行 npm test。",
            "Bilibili 演示视频仍需小组录制并以【武理26软工实践】作为标题前缀公开发布。",
        ]),
        ("七、AI 辅助开发说明", [
            "本项目使用 OpenAI Codex / GPT-5 系列能力辅助完成玩法设计、代码生成、测试设计、性能诊断、README 和报告初稿。AI 仅作为辅助工具，最终内容需小组成员审查、运行验证并按实际分工补充。",
        ]),
        ("八、小组分工记录", [
            "成员姓名、学号、角色、实际分工、Issue、PR 和提交记录需要小组按真实情况补充，不应由 AI 伪造。",
        ]),
        ("九、总结与后续工作", [
            "当前项目已经形成可运行、可测试、可展示的横版射击游戏。后续可继续补充排行榜、音效、更多 Boss、关卡配置、GitHub Issue 分工记录和演示视频链接。",
        ]),
    ]

    for title, paragraphs in sections:
        story.append(Paragraph(p(title), h1))
        for text in paragraphs:
            story.append(Paragraph(p(text), body))
        if title == "三、总体设计":
            story.append(make_pdf_table([
                ["模块", "文件", "职责"],
                ["入口层", "src/main.js", "输入、HUD、主循环"],
                ["规则层", "src/game/rules.js", "自动开火、连击、奖励、Boss 节奏"],
                ["状态层", "src/game/game.js", "实体、碰撞、掉落、BOMB、Boss"],
                ["渲染层", "src/game/render.js", "Canvas 绘制和粒子特效"],
                ["测试层", "tests/*.test.js", "规则和状态机测试"],
            ], body, [0.85 * inch, 1.55 * inch, 3.75 * inch]))
        if title == "八、小组分工记录":
            story.append(make_pdf_table([
                ["成员", "学号", "主要工作", "仓库证据"],
                ["待补充", "待补充", "按实际填写", "提交、Issue、PR"],
                ["待补充", "待补充", "待补充", "待补充"],
                ["待补充", "待补充", "待补充", "待补充"],
                ["待补充", "待补充", "待补充", "待补充"],
            ], body, [0.85 * inch, 0.9 * inch, 2.5 * inch, 1.6 * inch]))

    story.append(PageBreak())
    story.append(Paragraph("附录：运行与测试命令", h1))
    story.append(Paragraph("运行：npm start，访问 http://127.0.0.1:4173/", body))
    story.append(Paragraph("测试：npm test", body))

    doc = SimpleDocTemplate(PDF_OUT, pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch)
    doc.build(story)


def make_pdf_table(rows, body_style, widths):
    data = []
    for row in rows:
        data.append([Paragraph(p(str(cell)), body_style) for cell in row])
    table = Table(data, colWidths=widths, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DADCE0")),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table


if __name__ == "__main__":
    main()
